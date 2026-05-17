"""One-shot script to generate the binary test fixtures (PDF + DOCX).

Run once locally:  python tests/_make_fixtures.py

The generated files are committed to tests/fixtures/ so the pytest suite
has no runtime dependency on reportlab (used only here for PDF generation).
"""
from pathlib import Path

from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from docx import Document


FIX = Path(__file__).resolve().parent / "fixtures"
FIX.mkdir(parents=True, exist_ok=True)


def make_pdf():
    out = FIX / "sample.pdf"
    c = canvas.Canvas(str(out), pagesize=LETTER)
    c.setFont("Helvetica", 12)
    c.drawString(72, 720, "LexGuard test fixture.")
    c.drawString(72, 700, "This document confirms PDF extraction works.")
    c.drawString(72, 680, "Clause A: the signer agrees to test conditions.")
    c.showPage()
    c.save()
    print("wrote", out)


def make_docx():
    out = FIX / "sample.docx"
    doc = Document()
    doc.add_heading("LexGuard test fixture", level=1)
    doc.add_paragraph("This document confirms DOCX extraction works.")
    doc.add_paragraph("Clause A: the signer agrees to test conditions.")
    doc.save(out)
    print("wrote", out)


if __name__ == "__main__":
    make_pdf()
    make_docx()
