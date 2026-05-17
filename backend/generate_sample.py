"""One-shot script used to produce the bundled sample contract.
Run: python generate_sample.py
Output: ../frontend/public/sample-employment.docx
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path


TITLE = "EMPLOYMENT AGREEMENT"
INTRO = (
    "This Employment Agreement (\"Agreement\") is entered into between Nexora Systems Inc. "
    "(the \"Company\") and the undersigned individual (the \"Employee\"), effective as of the Employee's start date. "
    "By signing this Agreement, the Employee accepts all terms set forth herein."
)

CLAUSES = [
    ("1. Position and Duties",
     "The Company reserves the sole and absolute right to modify the Employee's role, title, "
     "responsibilities, reporting structure, and place of work at any time, with or without notice, "
     "and the Employee agrees that no such modification shall constitute grounds for resignation for "
     "cause or any claim against the Company."),

    ("2. Compensation",
     "The Company may, at its sole discretion, revise, defer, or withhold any portion of compensation, "
     "including accrued bonuses, where it determines that business conditions warrant. Bonuses, if any, "
     "are discretionary and are not earned until actually paid. The Employee waives any claim to bonus "
     "amounts not yet disbursed, including amounts accrued during a period of employment that subsequently ends."),

    ("3. Working Hours",
     "The Employee agrees to perform additional work as required, including evenings, weekends, and public "
     "holidays, without additional compensation. The Employee's salary is deemed to cover all hours worked "
     "regardless of quantity, and the Employee acknowledges that the role requires availability beyond "
     "standard hours."),

    ("4. Intellectual Property",
     "All inventions, ideas, works of authorship, designs, processes, and improvements conceived, developed, "
     "or reduced to practice by the Employee during the term of employment, whether or not related to the "
     "Company's business and whether or not created during working hours or using Company resources, shall "
     "be the sole and exclusive property of the Company. The Employee hereby irrevocably assigns all such "
     "rights to the Company."),

    ("5. Non-Competition",
     "For a period of twenty-four (24) months following the termination of employment for any reason, the "
     "Employee shall not, directly or indirectly, engage in, be employed by, consult for, or hold any "
     "ownership interest in any business that competes with the Company anywhere in the world. The Employee "
     "acknowledges that this restriction is reasonable in scope and duration."),

    ("6. Non-Solicitation",
     "For a period of thirty-six (36) months following termination, the Employee shall not solicit, hire, "
     "or attempt to hire any current or former employee, contractor, customer, prospect, or business partner "
     "of the Company. This restriction applies regardless of whether the Employee initiated the contact."),

    ("7. Confidentiality",
     "The Employee agrees to maintain in strict confidence, in perpetuity and without geographic limitation, "
     "all information of any kind learned during the course of employment, including information that is or "
     "may become publicly known. The Employee acknowledges that any breach shall entitle the Company to "
     "liquidated damages of not less than USD 250,000 per occurrence."),

    ("8. Termination",
     "The Company may terminate this Agreement at any time, with or without cause and without notice or "
     "severance. The Employee must provide ninety (90) days written notice prior to resignation, during "
     "which the Company may, at its discretion, place the Employee on garden leave or accelerate the "
     "termination date without further compensation."),

    ("9. Liability and Indemnification",
     "The Employee shall indemnify and hold the Company harmless from any losses, damages, claims, costs, "
     "or expenses (including legal fees) arising from the Employee's acts or omissions, regardless of "
     "whether such acts were within the scope of employment or authorized by the Company. The Employee's "
     "liability under this section shall not be capped."),

    ("10. Dispute Resolution",
     "Any dispute arising out of or relating to this Agreement shall be resolved by final and binding "
     "arbitration administered in Wilmington, Delaware, regardless of the Employee's place of residence "
     "or work. The Employee waives any right to participate in any class, collective, or representative "
     "action. The losing party shall bear all costs of arbitration including the prevailing party's "
     "attorney fees."),

    ("11. Modification",
     "The Company may amend, modify, or replace any term of this Agreement at any time by providing notice "
     "to the Employee. Continued employment after notice of any such change shall constitute the Employee's "
     "acceptance of the modified terms."),
]


def build():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading(TITLE, level=0)
    title.alignment = 1  # center

    doc.add_paragraph(INTRO)
    doc.add_paragraph()

    for heading, body in CLAUSES:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)

    doc.add_paragraph()
    doc.add_paragraph("By signing below, the Employee acknowledges that they have read, understood, "
                      "and agreed to all terms set forth in this Agreement.")
    doc.add_paragraph()
    doc.add_paragraph("Employee Signature: ____________________________   Date: ____________")
    doc.add_paragraph("For the Company: ____________________________   Date: ____________")

    out = Path(__file__).resolve().parent.parent / "frontend" / "public" / "sample-employment.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    build()
